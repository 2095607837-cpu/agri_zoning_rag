"""
农业区划数据解析器
解析 DOCX（技术规范）/ PDF（区划报告）/ XLSX（指标数据）/ CSV（台站数据）
输出统一 chunk JSON → data/chunks.json

统一架构: Parser → List[Block] → Heading Detect → Section → Chunk

用法:
  python3 step1_parse.py
  DATA_SRC=/path/to/source python3 step1_parse.py  # 指定数据源目录
"""

import json
import os
import re
import csv
import hashlib
from dataclasses import dataclass
from typing import Optional
from pathlib import Path
from collections import defaultdict

from table_linearizer import (
    linearize,
    _parse_pipe_table,
    _linearize_parsed,
    _build_context,
    _detect_table_schema,
    _build_caption,
)

BASE_DIR = Path(__file__).resolve().parent
DATA_SRC = Path(os.environ.get("DATA_SRC", "/Users/han/大模型自研代码/农业区划算法"))
OUTPUT = BASE_DIR / "data" / "chunks.json"

SKIP_DIRS = [
    "黑龙江农业气候资源普查和大豆区划规范-算法提交",
]
SKIP_NESTED = re.compile(r"05江西/05江西")
CSV_ENCODINGS = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]


# ═══════════════════════════════════════════════════════════
#  Unified Block & Heading Pipeline
# ═══════════════════════════════════════════════════════════

@dataclass
class Block:
    """统一文本块，Parser 无关。DOCX 带 style，PDF style 为 None。"""
    text: str
    style: Optional[str] = None    # DOCX: "Heading 1"/"Heading 2"; PDF: None
    page: Optional[int] = None     # PDF 页码
    source: str = ""               # 源文件名（用于调试）
    font_size: Optional[float] = None  # PDF dict mode: 字号(pt)
    bold: bool = False                 # PDF dict mode: 是否加粗
    x0: Optional[float] = None         # PDF dict mode: 左上角 x 坐标
    y0: Optional[float] = None         # PDF dict mode: 左上角 y 坐标
    is_title: bool = False             # PDF dict mode: 布局检测为标题


# 多 Regex 标题模式列表（按优先级从高到低）
_TIER2_PATTERNS = [
    (re.compile(r'^第[一二三四五六七八九十\d]+章[\s　]'),                                         'chapter'),   # 第6章 / 第一章
    (re.compile(r'^第[一二三四五六七八九十\d]+\s+章'),                                            'chapter_sp'), # 第6 章（数字章间有空格）
    (re.compile(r'^第[一二三四五六七八九十\d]+节[\s　]'),                                         'section'),   # 第2节
    (re.compile(r'^[一二三四五六七八九十]+[、．.\s]'),                                          'cn_num'),     # 一、二、
    (re.compile(r'^（[一二三四五六七八九十]+）'),                                               'cn_paren'),   # （一）（二）
    (re.compile(r'^\d+\.\d+\.\d+(?:\.\d+)?'),                                                  'num_dot3'),   # 5.2.4 / 2.3.1.1
    (re.compile(r'^\d+\.\d+'),                                                                  'num_dot2'),   # 5.2 / 6.1
    (re.compile(r'^\d+[\.\、]'),                                                                  'num_dot1'),   # 1. / 3、
    (re.compile(r'^（\d+）'),                                                                    'digit_paren'),# （1）（2）
    (re.compile(r'^[\(（]\d+[\)）][\s　]'),                                                     'digit_paren2'),# (1) / 1)
    (re.compile(r'^附录[一二三四五六七八九十A-Z]?'),                                            'appendix'),   # 附录A / 附录一
    (re.compile(r'^[A-Z]\.\s'),                                                                  'letter_dot'), # A. / B.
]


def _is_heading(text: str) -> bool:
    """多 Regex 列表逐个匹配，任意命中即为标题行。"""
    text = text.strip()
    if not text:
        return False
    for pattern, _ in _TIER2_PATTERNS:
        if pattern.match(text):
            return True
    return False


def _heading_level_from_style(style: Optional[str]) -> int:
    """从段落样式名提取标题级别，非标题返回 0。"""
    if not style:
        return 0
    m = re.search(r'(?:heading|标题|TOC)\s*(\d+)', style, re.IGNORECASE)
    if m:
        return int(m.group(1))
    if re.search(r'heading|标题', style, re.IGNORECASE):
        return 1
    return 0


# ── Section Builder ────────────────────────────────────

def _build_sections(blocks: list[Block]) -> list[dict]:
    """统一 Section 构建管道。
    Tier 1: 样式标题（DOCX 有 style 的 block）
    Tier 2: 多 Regex 标题识别（所有 block）
    Tier 3: 固定长度切分（兜底）
    """
    # Tier 1: 尝试样式标题
    style_sections = _build_style_sections(blocks)
    if len(style_sections) >= 2:
        return style_sections

    # Tier 2: 多 Regex 标题识别
    regex_sections = _build_regex_sections(blocks)
    if len(regex_sections) >= 2:
        return regex_sections

    # Tier 3: 固定长度切分（~800 字，句边界对齐）
    return _build_length_sections(blocks)


def _build_style_sections(blocks: list[Block]) -> list[dict]:
    """Tier 1: 基于 DOCX 样式标题构建 Section。"""
    h1_title = ""
    sections = []
    current = {"heading_path": [], "blocks": []}

    for b in blocks:
        level = _heading_level_from_style(b.style)
        text = b.text

        if level == 1:
            h1_title = text
        elif level == 2:
            if current["blocks"]:
                sections.append(current)
            path = [h1_title, text] if h1_title else [text]
            current = {"heading_path": path, "blocks": []}
        elif level >= 3:
            prefix = "#" * min(level, 4)
            current["blocks"].append(f"{prefix} {text}")
        else:
            current["blocks"].append(text)

    if current["blocks"]:
        sections.append(current)
    return sections


def _build_regex_sections(blocks: list[Block]) -> list[dict]:
    """Tier 2: 多 Regex 匹配标题行构建 Section。"""
    sections = []
    current = {"heading_path": [], "blocks": []}

    for b in blocks:
        text = b.text.strip()
        if not text:
            continue
        # dict 模式有布局信息时，仅信任 is_title（避免正则误匹配列表项）
        # text 回退模式无布局信息时，用正则 _is_heading()
        has_layout = b.font_size is not None
        is_split = b.is_title if has_layout else _is_heading(text)
        if is_split:
            if current["blocks"]:
                sections.append(current)
            current = {"heading_path": [text], "blocks": []}
        else:
            current["blocks"].append(text)

    if current["blocks"]:
        sections.append(current)
    return sections


def _build_length_sections(blocks: list[Block]) -> list[dict]:
    """Tier 3: 固定长度切分（兜底）。"""
    all_text = "\n".join(b.text for b in blocks if b.text.strip())
    parts = _split_by_sentence(all_text, 800)
    return [{"heading_path": [""], "blocks": [p]} for p in parts]


def _split_by_sentence(text: str, max_chars: int = 800) -> list[str]:
    """按句边界切分，每段 ≤ max_chars。"""
    sentences = re.split(r'(?<=[。！？\n])\s*', text)
    chunks = []
    current = ""
    for s in sentences:
        if not s.strip():
            continue
        if len(current) + len(s) > max_chars and current:
            chunks.append(current.strip())
            current = s
        else:
            current += s
    if current.strip():
        chunks.append(current.strip())
    return chunks or [text]


# ── Section Quality Filter (三层优先级架构) ──────────
# Layer 0: 结构识别 (Table / TOC / Fragment) → Drop 或 Group-Merge
# Layer 1: 语义判断 (完整句→Keep / PDF断句→Merge / 过渡句→Merge / 空标题→Merge)
# Layer 2: 长度兜底  (<80 且无完整句 → Merge)

# ── 预编译 Regex ─────────────────────────────────────────

_TOC_DOT_PAT = re.compile(r'[.…]{3,}\s*\d+\s*$')
_TOC_TAB_PAT = re.compile(r'\t{2,}\d+\s*$')
_TOC_SPACE_PAT = re.compile(r'\s{4,}\d+\s*$')
_TOC_MARKER_PAT = re.compile(r'###\s+\d[\d.]*\s+\S.*\d+')
_DEF_PAT = re.compile(r'是指|即|公式|定义')
_SENT_END = re.compile(r'[。！？]')
_LIST_PAT = re.compile(r'^\s*(?:[（\(]?\d+[）\)\.\、]|[一二三四五六七八九十]+[、．]|[-•·*])')
_DOMAIN_PAT = re.compile(r'区划|风险|作物|指标|气候|农业|气象|品种|种植|产量|品质|灾害|温度|降水|干旱|冷害|渍涝|霜冻|病虫害|土壤|光照|积温|日照|海拔|高原|地形|地势|地貌|丘陵|平原|山地|盆地|谷地|高程|纬度|经度')
TRANSITION_PATTERNS = ["包括以下", "主要包括", "如下", "具体如下", "分别为", "分为",
                       "以下几个方面", "以下方面", "主要措施", "对策建议"]
FOOTER_NOISE = re.compile(r'(区划报告|技术规范|初稿)\d*$|^\d+$|图\d+\.\d|表\d+\.\d')

# ── Helpers ──────────────────────────────────────────────

def _body_text(sec: dict) -> str:
    return "\n".join(sec.get("blocks", [])).strip()


def _is_mid_sentence_end(text: str) -> bool:
    """文本不以句子结束标点结尾 → 可能是断句。"""
    text = text.strip()
    if not text:
        return False
    return not text.endswith(('。', '！', '？', '；', '：', ':', ';'))


def _next_starts_continuation(next_sec: dict) -> bool:
    """下一个 section 开头是小写字母/数字/汉字（非大写开头）→ 续接上文。"""
    body = _body_text(next_sec)
    if not body:
        return True
    first = body[0]
    return (first.islower() or first.isdigit() or
            ('一' <= first <= '鿿'))


# ── Layer 0: 结构识别 ───────────────────────────────────

_TABLE_CAPTION_PAT = re.compile(r'^表\s*[0-9一二三四五六七八九十]+([-.．][0-9]+)*')
_UNIT_SYM_PAT = re.compile(
    r'℃|°C|mm|毫米|cm|厘米|米|km|'
    r'%|％|kg|g|t|hm²|亩|'
    r'≥|≤|>|<'
)


def _is_table_block(sec: dict) -> bool:
    """识别表格结构：pipe 硬通过 → 数字行占比 → 多特征打分（无 pipe 的 PDF 表格）。

    Stage 1 (硬规则，存量行为不变):
      - pipe > 3: markdown 表格（DOCX/XLSX 产出），强结构信号
      - 数字行 > 60%: fitz 合并成数字密集长行的 PDF 表格
    Stage 2 (多特征打分，补召回):
      - PDF 提取会把表格单元格拆成逐行短文本（0 pipe、数字行占比略低于 60%），
        Stage 1 漏判 → type=text → 跳过线性化 → 检索失败（如 Q_L07 陕西苹果区划指标表）。
      - 特征: 表题行前5行(+2), 数字行占比分段(≥0.6→+2, ≥0.5→+1),
        单位符号(+1), 连续数字行(+1), 阈值 4.
    """
    body = _body_text(sec)
    if not body:
        return False
    # Stage 1a: Pipe 数量 > 3（至少 2 列）
    if body.count('|') > 3:
        return True
    lines = [l.strip() for l in body.split("\n") if l.strip()]
    if len(lines) < 2:
        return False
    has_digit = [bool(re.search(r'\b\d+\.?\d*\b', l)) for l in lines]
    digit_lines = sum(has_digit)
    numeric_ratio = digit_lines / len(lines)
    # Stage 1b: 连续数字列（>60% 行含数字）
    if numeric_ratio > 0.6:
        return True
    # Stage 2: 多特征打分
    score = 0
    # 表题：前 5 个非空行（PDF 提取常在表题前插入页码）
    for line in lines[:5]:
        if _TABLE_CAPTION_PAT.match(line):
            score += 2
            break
    # 数字行占比：分段得分
    if numeric_ratio >= 0.60:
        score += 2
    elif numeric_ratio >= 0.50:
        score += 1
    if len(_UNIT_SYM_PAT.findall(body)) >= 2:
        score += 1
    # 连续数字行：正文极少出现数字-数字-数字的密集排列
    max_consec = consec = 0
    for d in has_digit:
        consec = consec + 1 if d else 0
        max_consec = max(max_consec, consec)
    if max_consec >= 5:
        score += 1
    return score >= 4


def _merge_table_block(sections: list[dict], start_idx: int) -> tuple[int, list[dict]]:
    """找出从 start_idx 开始的连续表格 section，合并所有 blocks 到第一个。"""
    # 先收集连续表格 section 的索引
    group = [start_idx]
    j = start_idx + 1
    while j < len(sections):
        if _is_table_block(sections[j]):
            group.append(j)
            j += 1
        else:
            break
    if len(group) == 1:
        return 0, []
    # 合并
    base = sections[group[0]]
    for idx in group[1:]:
        base["blocks"] = base["blocks"] + sections[idx]["blocks"]
    return len(group) - 1, [base]


def _is_toc_section(sec: dict) -> bool:
    """扩展 TOC 识别：点线目录 + tab/空格目录 + DOCX markdown 目录行。"""
    text = _body_text(sec)
    if not text:
        return False
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return False

    toc_count = 0
    for l in lines:
        if (_TOC_DOT_PAT.search(l) or _TOC_TAB_PAT.search(l) or
            _TOC_SPACE_PAT.search(l) or _TOC_MARKER_PAT.match(l)):
            toc_count += 1
    # 单行 TOC → drop; 多行且 >40% → drop
    return toc_count > 0 and (len(lines) == 1 or toc_count / len(lines) > 0.4)


# ── Layer 1: 语义判断 ───────────────────────────────────

def _is_complete_sentence(sec: dict) -> bool:
    """以句号/问号/感叹号结尾 → 完整语义句，即使短也应保留。"""
    body = _body_text(sec)
    return bool(body) and body.endswith(('。', '！', '？'))


def _is_pdf_fragment(sec: dict, next_sec: dict) -> bool:
    """PDF 截断片段：非句末结尾 + 80-250字 + 含领域词 + 下一 section 续接。"""
    body = _body_text(sec)
    if not (50 < len(body) < 250):
        return False
    if not _is_mid_sentence_end(body):
        return False
    if not _DOMAIN_PAT.search(body):
        return False
    return _next_starts_continuation(next_sec)


def _is_transition(sec: dict) -> bool:
    """识别过渡句：<120 字且含过渡关键词，或以冒号结尾。"""
    body = _body_text(sec)
    if not body or len(body) >= 120:
        return False
    if body.endswith((':', '：')):
        return True
    return any(p in body for p in TRANSITION_PATTERNS)


def _is_empty_title(sec: dict) -> bool:
    """标题后正文 <50 字 → 空标题。"""
    return len(_body_text(sec)) < 50


# ── Layer 2: 长度兜底 ───────────────────────────────────

def _is_short_fragment(sec: dict) -> bool:
    """极短且无完整句子 + 非定义/非列表 → 可合并。"""
    body = _body_text(sec)
    if not body or len(body) >= 80:
        return False
    if not _is_mid_sentence_end(body):
        return False
    if _DEF_PAT.search(body):
        return False
    if _LIST_PAT.match(body):
        return False
    return True


def _is_drop_junk(sec: dict) -> bool:
    """识别模板注释等无用内容。"""
    body = _body_text(sec)
    return any(m in body for m in ["不用删除", "标红内容为示例", "仅供参考"])


# ── 主过滤函数 ───────────────────────────────────────────

def _filter_sections(sections: list[dict]) -> list[dict]:
    results = []
    stats = {"dropped_toc": 0, "dropped_junk": 0,
             "merged_table": 0, "merged_transition": 0,
             "merged_empty": 0, "merged_broken": 0, "merged_short": 0,
             "kept_short_complete": 0}

    i = 0
    while i < len(sections):
        sec = sections[i]

        # ── Layer 0: 结构识别 ──

        # L0a: Drop junk (模板注释)
        if _is_drop_junk(sec):
            stats["dropped_junk"] += 1
            i += 1
            continue

        # L0b: Drop TOC
        if _is_toc_section(sec):
            stats["dropped_toc"] += 1
            i += 1
            continue

        # L0c: Table block — 识别并合并连续表格 section
        if _is_table_block(sec):
            merged_n, table_results = _merge_table_block(sections, i)
            if merged_n > 0:
                stats["merged_table"] += merged_n
                results.append(table_results[0])
                i += merged_n + 1
                continue

        # ── Layer 1: 语义判断 ──
        # 注意顺序：先清理垃圾（空标题/过渡句），再判断保留/断句合并

        # L1a: 空标题合并（<50字正文）— 最先，清理无内容标题
        if _is_empty_title(sec) and i + 1 < len(sections):
            sections[i + 1]["blocks"] = sec["blocks"] + sections[i + 1]["blocks"]
            stats["merged_empty"] += 1
            i += 2
            continue

        # L1b: 过渡句合并（<120字 + 过渡关键词/冒号结尾）
        if _is_transition(sec) and i + 1 < len(sections):
            sections[i + 1]["blocks"] = sec["blocks"] + sections[i + 1]["blocks"]
            stats["merged_transition"] += 1
            i += 2
            continue

        # L1c: 完整语义句 (<200字但以。！？结尾) → 保留
        if len(_body_text(sec)) < 200 and _is_complete_sentence(sec):
            stats["kept_short_complete"] += 1
            results.append(sec)
            i += 1
            continue

        # L1d: PDF 断句补全 (80-250字, 无句末标点, 含领域词)
        if (i + 1 < len(sections) and
            _is_pdf_fragment(sec, sections[i + 1])):
            sections[i + 1]["blocks"] = sec["blocks"] + sections[i + 1]["blocks"]
            stats["merged_broken"] += 1
            i += 2
            continue

        # ── Layer 2: 长度兜底 ──

        if _is_short_fragment(sec) and i + 1 < len(sections):
            sections[i + 1]["blocks"] = sec["blocks"] + sections[i + 1]["blocks"]
            stats["merged_short"] += 1
            i += 2
            continue

        results.append(sec)
        i += 1

    # 打印统计
    total_dropped = stats["dropped_toc"] + stats["dropped_junk"]
    total_merged = (stats["merged_table"] + stats["merged_transition"] +
                    stats["merged_empty"] + stats["merged_broken"] +
                    stats["merged_short"])
    if total_dropped or total_merged or stats["kept_short_complete"]:
        parts = [f"{len(sections)} -> {len(results)}"]
        if total_dropped:
            parts.append(f"丢弃{total_dropped}")
        if total_merged:
            parts.append(f"合并{total_merged}")
        print(f"    [Section 过滤] {' | '.join(parts)}")
        detail_parts = []
        if stats["dropped_toc"]:
            detail_parts.append(f"TOC={stats['dropped_toc']}")
        if stats["dropped_junk"]:
            detail_parts.append(f"模板注释={stats['dropped_junk']}")
        if stats["merged_table"]:
            detail_parts.append(f"表格={stats['merged_table']}")
        if stats["merged_broken"]:
            detail_parts.append(f"PDF断句={stats['merged_broken']}")
        if stats["merged_transition"]:
            detail_parts.append(f"过渡句={stats['merged_transition']}")
        if stats["merged_empty"]:
            detail_parts.append(f"空标题={stats['merged_empty']}")
        if stats["merged_short"]:
            detail_parts.append(f"短片段={stats['merged_short']}")
        if stats["kept_short_complete"]:
            detail_parts.append(f"保留短完整句={stats['kept_short_complete']}")
        if detail_parts:
            print(f"           [{', '.join(detail_parts)}]")

    return results


# ── Chunk 生成 ─────────────────────────────────────────

def _sections_to_chunks(sections: list[dict], doc_stem: str, fname: str,
                         province: str, crop: str, zoning: str,
                         source_type: str = "technical_spec",
                         layout_mode: str = "none") -> list[dict]:
    """将 Section 列表转为最终 chunk 列表。每个 Section 生成一个 chunk。"""
    chunks = []
    for si, sec in enumerate(sections):
        full_text = "\n\n".join(sec.get("blocks", [])).strip()
        if len(full_text) < 20:
            continue

        heading_path = sec.get("heading_path") or [doc_stem]
        section_id = f"{doc_stem}_sec_{si}"

        # 判断 section 内容类型
        if _is_table_block(sec):
            sec_type = "table"
        elif _is_heading(full_text[:60]) and len(full_text) < 60:
            sec_type = "heading"
        else:
            sec_type = "text"

        chunks.append({
            "id": f"{doc_stem}_s{si}",
            "content": full_text,
            "source_id": section_id,
            "chunk_version": 1,
            "metadata": {
                "doc_id": doc_stem,
                "section_id": section_id,
                "heading_path": heading_path,
                "heading_level": len(heading_path),
                "source_type": source_type,
                "source_file": fname,
                "province": province,
                "crop": crop,
                "zoning_type": zoning,
                "section_title": heading_path[-1] if heading_path else "",
                "type": sec_type,
                "layout_mode": layout_mode,
            },
        })
    return chunks


# ═══════════════════════════════════════════════════════════
#  DOCX Parser
# ═══════════════════════════════════════════════════════════

def parse_docx(filepath: str) -> list[dict]:
    """解析 DOCX：提取段落 → Block 列表 → 统一 Section 管道。"""
    try:
        import docx
    except ImportError:
        print(f"  [!] python-docx 未安装，跳过: {filepath}")
        return []

    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"  [!] 无法打开 {filepath}: {e}")
        return []

    fname = Path(filepath).name
    doc_stem = Path(filepath).stem
    province, crop, zoning = _infer_meta(fname)

    # 提取段落 → Block 列表
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        blocks.append(Block(text=text, style=style, source=fname))

    # 统一 Section 管道
    sections = _build_sections(blocks)
    sections = _filter_sections(sections)
    chunks = _sections_to_chunks(sections, doc_stem, fname, province, crop, zoning,
                                  source_type="technical_spec", layout_mode="docx_style")

    # 分配表格到最近的 section
    table_chunks = _extract_tables(doc, doc_stem, fname, province, crop, zoning)
    chunks.extend(table_chunks)

    return chunks


def _extract_tables(doc, doc_stem: str, fname: str,
                    province: str, crop: str, zoning: str) -> list[dict]:
    """提取 DOCX 表格为独立 chunk。"""
    chunks = []
    for i, table in enumerate(doc.tables):
        table_text = _table_to_markdown(table)
        if len(table_text) < 30:
            continue
        section_id = f"{doc_stem}_table_{i}"
        chunks.append({
            "id": f"{doc_stem}_t{i}",
            "content": table_text,
            "source_id": section_id,
            "chunk_version": 1,
            "metadata": {
                "doc_id": doc_stem,
                "section_id": section_id,
                "heading_path": [f"表格{i+1}"],
                "heading_level": 1,
                "source_type": "technical_spec",
                "source_file": fname,
                "province": province,
                "crop": crop,
                "zoning_type": zoning,
                "section_title": f"表格{i+1}",
                "type": "table",
                "layout_mode": "docx_style",
            },
        })
    return chunks


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    if len(rows) > 1:
        header_sep = "|" + "|".join(["---"] * len(rows[0].split("|"))) + "|"
        rows.insert(1, header_sep)
    return "\n".join(rows)


# ═══════════════════════════════════════════════════════════
#  PDF Parser (全文拼接 + 多 Regex Heading)
# ═══════════════════════════════════════════════════════════

def parse_pdf(filepath: str) -> list[dict]:
    """解析 PDF：优先使用 fitz dict 模式提取排版信息，失败则回退全文拼接。"""
    try:
        import fitz
    except ImportError:
        print(f"  [!] pymupdf 未安装，跳过: {filepath}")
        return []

    fname = Path(filepath).name
    doc_stem = Path(filepath).stem
    province, crop, zoning = _infer_meta(fname)

    try:
        doc = fitz.open(filepath)
    except Exception as e:
        print(f"  [!] 无法打开 PDF {filepath}: {e}")
        return []

    num_pages = len(doc)
    layout_mode = "text_fallback"
    all_blocks = []

    # ── 尝试 dict 模式 ──
    try:
        layout_blocks, ok = _extract_layout_blocks(doc)
        if ok and len(layout_blocks) >= 10:
            # 按页跳过 TOC
            page_texts = defaultdict(list)
            for b in layout_blocks:
                if b.page is not None:
                    page_texts[b.page].append(b.text)
            toc_pages = set()
            for pn, texts in page_texts.items():
                joined = "\n".join(texts)
                if len(joined) < 50 or _is_toc_page(joined):
                    toc_pages.add(pn)
            layout_blocks = [b for b in layout_blocks if b.page not in toc_pages]

            if layout_blocks:
                layout_blocks = _dedup_headers_footers(layout_blocks, num_pages)
                _detect_titles_by_layout(layout_blocks)
                all_blocks = layout_blocks
                layout_mode = "dict"
                title_n = sum(1 for b in all_blocks if b.is_title)
                print(f"    [PDF Layout] dict 模式, {len(all_blocks)} blocks, "
                      f"{title_n} 标题")
    except Exception as e:
        print(f"    [PDF Layout] dict 异常: {e}, 回退 text 模式")

    # ── 回退 text 模式 ──
    if not all_blocks:
        all_text_parts = []
        for page_num in range(num_pages):
            page = doc[page_num]
            text = page.get_text().strip()
            if len(text) < 50:
                continue
            if _is_toc_page(text):
                continue
            all_text_parts.append(text)

        if not all_text_parts:
            doc.close()
            return []

        full_text = "\n".join(all_text_parts)
        full_text = _clean_pdf_text(full_text)
        all_blocks = _lines_to_blocks(full_text, fname)

    doc.close()

    if not all_blocks:
        return []

    sections = _build_sections(all_blocks)
    sections = _filter_sections(sections)
    return _sections_to_chunks(sections, doc_stem, fname, province, crop, zoning,
                                source_type="report", layout_mode=layout_mode)


def _is_toc_page(text: str) -> bool:
    """检测是否为目录页。"""
    lines = text.split("\n")
    dot_lines = 0
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 目录行特征: 标题 ... 页码 或 罗马数字章节头
        if "...." in line or "……" in line:
            dot_lines += 1
        if re.match(r'^[IVX]+\s*$', line):
            return True
    # 超过 40% 的行含省略号 → 目录页
    return dot_lines > len(lines) * 0.4 if lines else False


def _clean_pdf_text(text: str) -> str:
    """清洗 PDF 全文：去页眉页脚，保留段落结构（不合并行内换行）。"""
    # 移除罗马数字独立行（页眉页脚）
    text = re.sub(r'\n[IVX]+\n', '\n', text)
    # 移除纯数字独立行（页码）
    text = re.sub(r'\n\d{1,3}\n', '\n', text)
    # 压缩三行以上空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _lines_to_blocks(text: str, source: str) -> list[Block]:
    """将文本按段落边界转为 Block 列表，标题行独立，正文行合并。
    策略：
      1. 按双换行切段落
      2. 段落内按单换行拆行
      3. 标题行 → 独立 Block
      4. 连续非标题行 → 合并为一个 Block
    """
    blocks = []
    paragraphs = text.split("\n\n")

    for para in paragraphs:
        lines = para.strip().split("\n")
        if not lines:
            continue

        body_buf = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                body_buf.append("")
                continue

            if _is_heading(stripped):
                # 先刷出积攒的正文
                if body_buf:
                    merged = "".join(body_buf).strip()
                    if len(merged) >= 20:
                        blocks.append(Block(text=merged, source=source))
                    body_buf = []
                # 标题独立成块
                blocks.append(Block(text=stripped, source=source))
            else:
                body_buf.append(stripped)

        if body_buf:
            merged = "".join(body_buf).strip()
            if len(merged) >= 20:
                blocks.append(Block(text=merged, source=source))

    return blocks


# ── PDF Layout 解析（Phase 1: fitz dict 模式）─────────────

def _extract_layout_blocks(doc) -> tuple:
    """使用 page.get_text("dict") 提取带排版信息的 Block 列表。

    Returns:
        (blocks: list[Block], success: bool)
        每个 line 变成一个 Block，合并 line 内多个 span。
    """
    blocks = []
    for page_num in range(len(doc)):
        try:
            page = doc[page_num]
            pagedict = page.get_text("dict")
            for block in pagedict.get("blocks", []):
                if block.get("type") != 0:  # 跳过图片块
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    text_parts = []
                    max_font = 0.0
                    any_bold = False
                    x0 = spans[0].get("bbox", [0, 0, 0, 0])[0]
                    y0 = spans[0].get("bbox", [0, 0, 0, 0])[1]
                    for sp in spans:
                        text_parts.append(sp.get("text", ""))
                        max_font = max(max_font, sp.get("size", 0))
                        if sp.get("flags", 0) & 16:  # bit 4 = bold
                            any_bold = True
                    text = "".join(text_parts).strip()
                    if text:
                        blocks.append(Block(
                            text=text,
                            page=page_num,
                            font_size=max_font if max_font > 0 else None,
                            bold=any_bold,
                            x0=x0,
                            y0=y0,
                        ))
        except Exception:
            return [], False
    return blocks, len(blocks) > 0


def _detect_titles_by_layout(blocks: list[Block]):
    """基于字号 P70/P90 双阈值 + 加粗 + 长度 检测标题。

    评分规则:
      - font_size >= P90 → 0.5 (章节标题级)
      - font_size >= P70 → 0.3 (小节标题级)
      - bold           → 0.1
      - len(text) < 30 → 0.2
      - score > 0.6    → is_title = True
    """
    font_sizes = sorted(
        [b.font_size for b in blocks if b.font_size and b.font_size > 0]
    )
    if len(font_sizes) < 5:
        return

    n = len(font_sizes)
    p70 = font_sizes[int(n * 0.7)]
    p90 = font_sizes[int(n * 0.9)]

    for b in blocks:
        score = 0.0
        if b.font_size and b.font_size > 0:
            if b.font_size >= p90:
                score += 0.5
            elif b.font_size >= p70:
                score += 0.3
        if b.bold:
            score += 0.1
        if len(b.text.strip()) < 30:
            score += 0.2
        if score > 0.6:
            b.is_title = True


def _dedup_headers_footers(blocks: list[Block], num_pages: int) -> list[Block]:
    """移除跨页重复的页眉/页脚。

    策略:
      1. 每页按 y0 排序，取前 2 / 后 2 个 Block 为候选
      2. 候选项在 >40% 页面出现 → 标记为噪声
      3. 仅移除处于页眉/页脚位置的噪声 Block（不影响正文中的同名文本）
    """
    from collections import defaultdict

    page_blocks = defaultdict(list)
    for b in blocks:
        if b.page is not None:
            page_blocks[b.page].append(b)

    if len(page_blocks) < 3:
        return blocks

    # 每页按 y0 排序，收集首尾候选
    candidates = defaultdict(int)  # hash → 出现页数
    page_sorted = {}
    for pg, pblocks in page_blocks.items():
        sorted_b = sorted(
            [b for b in pblocks if b.y0 is not None],
            key=lambda b: b.y0,
        )
        if not sorted_b:
            continue
        page_sorted[pg] = sorted_b
        seen = set()
        for b in sorted_b[:2]:
            t = b.text.strip()
            if len(t) < 100:
                h = hashlib.md5(t[:100].encode()).hexdigest()
                if h not in seen:
                    candidates[h] += 1
                    seen.add(h)
        for b in sorted_b[-2:]:
            t = b.text.strip()
            if len(t) < 100:
                h = hashlib.md5(t[:100].encode()).hexdigest()
                if h not in seen:
                    candidates[h] += 1
                    seen.add(h)

    threshold = max(3, num_pages * 0.4)
    noise_hashes = {h for h, c in candidates.items() if c > threshold}

    if not noise_hashes:
        return blocks

    # 过滤：仅移除处于页眉/页脚位置的匹配 Block
    filtered = []
    for b in blocks:
        if b.page is not None and b.page in page_sorted:
            sorted_b = page_sorted[b.page]
            top_ids = {id(x) for x in sorted_b[:2]}
            bot_ids = {id(x) for x in sorted_b[-2:]}
            if id(b) in top_ids or id(b) in bot_ids:
                h = hashlib.md5(b.text.strip()[:100].encode()).hexdigest()
                if h in noise_hashes:
                    continue
        filtered.append(b)

    dedup_n = len(blocks) - len(filtered)
    if dedup_n:
        print(f"    [PDF Layout] 移除页眉页脚 {dedup_n} 行 (阈值>{threshold:.0f}页)")
    return filtered


# ═══════════════════════════════════════════════════════════
#  XLSX Parser（不变）
# ═══════════════════════════════════════════════════════════

def parse_xlsx(filepath: str) -> list[dict]:
    try:
        import openpyxl
    except ImportError:
        print(f"  [!] openpyxl 未安装，跳过: {filepath}")
        return []

    chunks = []
    fname = Path(filepath).name
    doc_stem = Path(filepath).stem
    province, crop, zoning = _infer_meta(fname)

    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
    except Exception as e:
        print(f"  [!] 无法打开 XLSX {filepath}: {e}")
        return []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [str(h) if h else "" for h in rows[0]]
        data_rows = rows[1:]
        data_rows = [r for r in data_rows if any(c is not None for c in r)]

        if not data_rows:
            continue

        text_parts = [
            f"表格: {sheet_name}",
            f"列名: {', '.join(h for h in headers if h)}",
            f"数据行数: {len(data_rows)}",
            "",
        ]

        numeric_cols = {}
        for ci, h in enumerate(headers):
            vals = []
            for r in data_rows:
                v = r[ci] if ci < len(r) else None
                if isinstance(v, (int, float)):
                    vals.append(v)
            if vals:
                numeric_cols[h] = vals

        if numeric_cols:
            text_parts.append("数值列统计:")
            for col_name, vals in numeric_cols.items():
                text_parts.append(
                    f"  {col_name}: 最小值={min(vals):.4f}, 最大值={max(vals):.4f}, "
                    f"平均值={sum(vals)/len(vals):.4f}"
                )

        text_parts.append("\n前5行数据示例:")
        for ri, row in enumerate(data_rows[:5]):
            row_str = " | ".join(str(c) if c is not None else "" for c in row)
            text_parts.append(f"  行{ri + 2}: {row_str}")

        full_text = "\n".join(text_parts)
        section_id = f"{doc_stem}_{sheet_name}"
        chunks.append({
            "id": f"{doc_stem}_{sheet_name}",
            "content": full_text,
            "source_id": section_id,
            "chunk_version": 1,
            "metadata": {
                "doc_id": doc_stem,
                "section_id": section_id,
                "heading_path": [sheet_name],
                "heading_level": 1,
                "source_type": "data_table",
                "source_file": fname,
                "province": province,
                "crop": crop,
                "zoning_type": zoning,
                "sheet_name": sheet_name,
                "section_title": sheet_name,
                "type": "table",
                "layout_mode": "xlsx",
            },
        })

    wb.close()
    return chunks


# ═══════════════════════════════════════════════════════════
#  CSV Parser（不变）
# ═══════════════════════════════════════════════════════════

def parse_csv(filepath: str) -> list[dict]:
    fname = Path(filepath).name
    doc_stem = Path(filepath).stem
    province, crop, zoning = _infer_meta(fname)

    rows = None
    for enc in CSV_ENCODINGS:
        try:
            with open(filepath, "r", encoding=enc) as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception:
            break

    if rows is None:
        print(f"  [!] 无法解码 CSV {filepath}，已尝试编码: {CSV_ENCODINGS}")
        return []

    if not rows:
        return []

    headers = list(rows[0].keys())
    field_info = "\n".join([f"- {h}" for h in headers])

    numeric_stats = []
    for h in headers:
        try:
            vals = [float(r[h]) for r in rows if r[h] and r[h].strip()]
            if vals:
                numeric_stats.append(
                    f"  {h}: 范围 [{min(vals):.2f}, {max(vals):.2f}], "
                    f"均值 {sum(vals)/len(vals):.2f}, 共 {len(vals)} 条"
                )
        except ValueError:
            pass

    text_parts = [
        f"CSV 数据文件: {fname}",
        f"总行数: {len(rows)}",
        f"\n字段列表:\n{field_info}",
    ]
    if numeric_stats:
        text_parts.append(f"\n数值字段统计:")
        text_parts.extend(numeric_stats)

    text_parts.append(f"\n前5行数据示例:")
    for ri, row in enumerate(rows[:5]):
        items = [f"{k}={v}" for k, v in row.items()]
        text_parts.append(f"  [{ri + 1}] {', '.join(items)}")

    return [{
        "id": doc_stem,
        "content": "\n".join(text_parts),
        "source_id": doc_stem,
        "chunk_version": 1,
        "metadata": {
            "doc_id": doc_stem,
            "section_id": doc_stem,
            "heading_path": [fname],
            "heading_level": 1,
            "source_type": "data_table",
            "source_file": fname,
            "province": province,
            "crop": crop,
            "zoning_type": zoning,
            "section_title": fname,
            "type": "table",
            "layout_mode": "csv",
        },
    }]


# ═══════════════════════════════════════════════════════════
#  元数据推断（不变）
# ═══════════════════════════════════════════════════════════

PROVINCE_MAP = {
    "内蒙古": "内蒙古", "黑龙江": "黑龙江", "河南": "河南",
    "江西": "江西", "陕西": "陕西", "新疆": "新疆",
    "辽宁": "辽宁", "气科院": "全国",
}

CROP_MAP = {
    "大豆": "大豆", "小麦": "冬小麦", "冬小麦": "冬小麦",
    "柑橘": "柑橘", "苹果": "苹果",
    "猕猴桃": "猕猴桃", "桃": "桃", "梨": "梨",
    "品质": None, "病害": None, "区划": None,
}

ZONING_MAP = {
    "产量": "产量区划", "品质": "品质区划", "种植": "种植区划",
    "干旱": "干旱风险区划", "冷害": "冷害风险区划",
    "渍涝": "渍涝风险区划", "霜冻": "霜冻风险区划",
    "病虫害": "病虫害风险区划", "食心虫": "病虫害风险区划",
    "气候区划": "气候区划", "气候风险": "气候风险区划",
    "算法": "算法说明", "技术规范": "技术规范",
    "普查": "资源普查", "指标": "区划指标",
}


def _infer_meta(filename: str):
    province, crop, zoning = "", "", ""
    for key, val in PROVINCE_MAP.items():
        if key in filename:
            province = val
            break
    for key, val in CROP_MAP.items():
        if key in filename:
            crop = val or crop
            break
    for key, val in ZONING_MAP.items():
        if key in filename:
            zoning = val
            break
    if not zoning:
        zoning = "农业气候区划"
    return province, crop, zoning


# ═══════════════════════════════════════════════════════════
#  质量标记 & 去重（不变）
# ═══════════════════════════════════════════════════════════

def _tag_chunk_quality(chunks: list[dict]) -> list[dict]:
    low_patterns = [
        (r'^[#\s]*(目\s*录|Contents|Table of Contents)', '目录'),
        (r'^[#\s]*(参考文献|References|参考书目)', '参考文献'),
    ]
    for c in chunks:
        content = c['content'].strip()
        reasons = []
        if len(content) < 50:
            reasons.append('极短(<50字)')
        for pat, label in low_patterns:
            if re.match(pat, content):
                reasons.append(label)
        c['quality'] = 'low' if reasons else 'high'
        c['quality_reason'] = '; '.join(reasons) if reasons else ''
    low_count = sum(1 for c in chunks if c['quality'] == 'low')
    print(f"\n  质量标记: high={len(chunks)-low_count} low={low_count}")
    return chunks


def _dedup_documents(chunks: list[dict]) -> list[dict]:
    groups = defaultdict(lambda: {'docx': defaultdict(list), 'pdf': defaultdict(list)})
    for i, c in enumerate(chunks):
        meta = c['metadata']
        key = (meta.get('province', ''), meta.get('crop', ''), meta.get('zoning_type', ''))
        src = meta.get('source_file', '')
        if src.endswith('.docx'):
            groups[key]['docx'][src].append(i)
        elif src.endswith('.pdf'):
            groups[key]['pdf'][src].append(i)

    excluded = set()
    for group_key, data in groups.items():
        if not data['docx'] or not data['pdf']:
            continue
        for docx_src, docx_idxs in data['docx'].items():
            for pdf_src, pdf_idxs in data['pdf'].items():
                docx_total = sum(len(chunks[i]['content']) for i in docx_idxs)
                pdf_total = sum(len(chunks[i]['content']) for i in pdf_idxs)
                ratio = max(docx_total, pdf_total) / min(docx_total, pdf_total)
                if ratio > 1.15:
                    if docx_total > pdf_total:
                        excluded.update(pdf_idxs)
                        print(f"  文档去重: 保留 DOCX ({docx_total}字) 排除 PDF ({pdf_total}字) [{group_key}]")
                    else:
                        excluded.update(docx_idxs)
                        print(f"  文档去重: 保留 PDF ({pdf_total}字) 排除 DOCX ({docx_total}字) [{group_key}]")
                else:
                    print(f"  文档去重: 内容量相当(ratio={ratio:.2f}) 都保留 [{group_key}]")

    if excluded:
        for i in excluded:
            chunks[i]['excluded'] = True
        print(f"\n  文档去重: 排除 {len(excluded)} 个 chunk")
    return chunks


# ═══════════════════════════════════════════════════════════
#  主流程
# ═══════════════════════════════════════════════════════════

def parse_all(src_dir: str) -> list[dict]:
    all_chunks = []
    src_path = Path(src_dir)

    for f in sorted(src_path.rglob("*")):
        if not f.is_file():
            continue

        fstr = str(f)
        skip = False
        for sd in SKIP_DIRS:
            if sd in fstr:
                skip = True
                break
        if SKIP_NESTED.search(fstr):
            skip = True
        if skip:
            continue

        suffix = f.suffix.lower()
        fname = f.name

        if suffix == ".docx":
            print(f"  [DOCX] {fname}")
            chunks = parse_docx(str(f))
        elif suffix == ".pdf":
            size_mb = f.stat().st_size / 1024 / 1024
            if size_mb > 50:
                print(f"  [PDF] {fname} ({size_mb:.0f}MB) — 跳过（过大）")
                continue
            print(f"  [PDF] {fname} ({size_mb:.1f}MB)")
            chunks = parse_pdf(str(f))
        elif suffix == ".xlsx":
            print(f"  [XLSX] {fname}")
            chunks = parse_xlsx(str(f))
        elif suffix == ".csv":
            print(f"  [CSV] {fname}")
            chunks = parse_csv(str(f))
        else:
            continue

        all_chunks.extend(chunks)
        if chunks:
            print(f"         → {len(chunks)} chunks")

    return all_chunks


def _linearize_and_split_tables(chunks: list[dict]) -> list[dict]:
    """对 table 类型 chunk 做线性化，并按 ≤800 字拆分为子 chunk（行窗口）。"""
    expanded = []
    table_count = 0
    linearized_count = 0
    split_count = 0

    for c in chunks:
        if c["metadata"].get("type") != "table":
            expanded.append(c)
            continue

        table_count += 1
        content = c["content"]
        metadata = c["metadata"]

        # 尝试线性化
        result = linearize(content, metadata)
        if result is None:
            expanded.append(c)
            continue

        linearized_count += 1

        # 按 ≤800 字拆分，保留行边界（以 。 为分隔）
        if len(result) <= 800:
            c["content"] = result
            expanded.append(c)
            continue

        # 分离前缀（上下文 + caption）和数据行
        first_period = result.find("。")
        if first_period < 0:
            c["content"] = result
            expanded.append(c)
            continue

        prefix = result[:first_period + 1]
        data = result[first_period + 1:]

        # 数据行按 。 拆分
        data_sentences = [s.strip() + "。" for s in data.split("。") if s.strip()]

        # 按 800 字分组
        sub_chunks = []
        current = prefix
        for s in data_sentences:
            if len(current) + len(s) > 800 and current != prefix:
                sub_chunks.append(current)
                current = prefix + s
            else:
                current += s

        if current != prefix:
            sub_chunks.append(current)

        if len(sub_chunks) <= 1:
            c["content"] = result
            expanded.append(c)
            continue

        split_count += 1
        base_id = c["id"]
        base_meta = dict(metadata)
        for i, sc in enumerate(sub_chunks):
            new_meta = dict(base_meta)
            new_meta["_table_chunk_index"] = i
            new_meta["_table_chunk_count"] = len(sub_chunks)
            expanded.append({
                "id": f"{base_id}_t{i}" if i > 0 else base_id,
                "content": sc,
                "source_id": c.get("source_id", c["metadata"].get("section_id", base_id)),
                "chunk_version": 1,
                "metadata": new_meta,
            })

    if linearized_count:
        print(f"\n  表格: {table_count} → 线性化 {linearized_count} → 拆分 {split_count}")
    return expanded


def main():
    print(f"数据源: {DATA_SRC}")
    print(f"输出: {OUTPUT}")
    print()

    raw_chunks = parse_all(str(DATA_SRC))

    # 第1轮：按 id 去重
    seen_ids = set()
    id_deduped = []
    for c in raw_chunks:
        if c["id"] not in seen_ids:
            seen_ids.add(c["id"])
            id_deduped.append(c)

    # 第2轮：按内容哈希去重
    seen_hashes = set()
    deduped = []
    for c in id_deduped:
        content_hash = hashlib.md5(c["content"][:200].encode()).hexdigest()
        if content_hash not in seen_hashes:
            seen_hashes.add(content_hash)
            c["content_hash"] = content_hash
            deduped.append(c)

    dup_count = len(raw_chunks) - len(id_deduped)
    content_dup = len(id_deduped) - len(deduped)
    if dup_count:
        print(f"\n  已去除 {dup_count} 个 ID 重复 chunk")
    if content_dup:
        print(f"  已去除 {content_dup} 个内容重复 chunk")

    deduped = _dedup_documents(deduped)
    deduped = _tag_chunk_quality(deduped)
    deduped = _linearize_and_split_tables(deduped)

    source_types = {}
    provinces = {}
    for c in deduped:
        st = c["metadata"].get("source_type", "unknown")
        source_types[st] = source_types.get(st, 0) + 1
        p = c["metadata"].get("province", "unknown")
        provinces[p] = provinces.get(p, 0) + 1

    print(f"\n{'=' * 60}")
    print(f"  解析完成")
    print(f"{'=' * 60}")
    print(f"  总 chunks: {len(deduped)}")
    print(f"  按类型: {json.dumps(source_types, ensure_ascii=False)}")
    print(f"  按省份: {json.dumps(provinces, ensure_ascii=False)}")

    avg_len = sum(len(c["content"]) for c in deduped) / len(deduped) if deduped else 0
    print(f"  平均 chunk 长度: {avg_len:.0f} 字符")

    os.makedirs(OUTPUT.parent, exist_ok=True)
    with open(OUTPUT, "w", encoding="utf-8") as f:
        json.dump(deduped, f, ensure_ascii=False, indent=2)

    print(f"\n  已保存到: {OUTPUT}")


if __name__ == "__main__":
    main()
